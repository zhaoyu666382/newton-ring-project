# -*- coding: utf-8 -*-
"""
Generate a Word report (.docx) with:
- Basic info
- Data table (n, r, r^2)
- Figures (overlay, radial profile, r^2-n fit)
- Results and conclusion

Font requirements:
- Use SimSun (宋体) for East Asian text (set via run._element.rPr.rFonts)
"""

from __future__ import annotations

from typing import Dict, Any
from pathlib import Path
import os
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def _set_run_font(run, font_name: str = "SimSun", font_size_pt: int = 12, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size_pt)
    run.bold = bold


def _set_paragraph_font(paragraph, font_name: str = "SimSun", font_size_pt: int = 12):
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, font_size_pt=font_size_pt)


def _make_fit_plot(fit: Dict[str, Any], out_path: str):
    n = np.asarray(fit["n"], dtype=float)
    r2 = np.asarray(fit["r2_mm2"], dtype=float)
    slope = float(fit["slope"])
    intercept = float(fit["intercept"])

    plt.figure(figsize=(8, 5), dpi=300)
    plt.scatter(n, r2, s=45, label="data")
    plt.plot(n, slope * n + intercept, linewidth=2, label=f"fit: r^2={slope:.4f} n + {intercept:.4f}")
    plt.xlabel("Ring index n")
    plt.ylabel("r^2 (mm^2)")
    plt.title("Newton rings: r^2 vs n")
    plt.grid(True, alpha=0.3)
    plt.legend()
    plt.tight_layout()
    os.makedirs(Path(out_path).parent, exist_ok=True)
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()


def generate_word_report(
    outdir: str,
    basename: str,
    det: Dict[str, Any],
    fit: Dict[str, Any],
    err=None,
    report_cfg: Dict[str, Any] | None = None,
    logger=None,
) -> str:
    report_cfg = report_cfg or {}
    outdir = Path(outdir)
    rep_dir = outdir / "reports"
    fig_dir = outdir / "figures"
    rep_dir.mkdir(parents=True, exist_ok=True)
    fig_dir.mkdir(parents=True, exist_ok=True)

    fit_fig = str(fig_dir / f"{basename}_fit_r2_n.png")
    _make_fit_plot(fit, fit_fig)

    doc = Document()

    # =====================
    # 标题
    # =====================
    title = doc.add_paragraph()
    run = title.add_run("牛顿环实验报告")
    _set_run_font(run, font_size_pt=18, bold=True)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    info = doc.add_paragraph()
    run = info.add_run(
        f"样本名称：{basename}    "
        f"圆心识别方法：{det['center']['method']}    "
        f"曲率半径 R = {fit['R_mm']:.2f} mm"
    )
    _set_run_font(run, font_size_pt=10)
    info.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()

    # =====================
    # 一、实验目的
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("一、实验目的"), font_size_pt=14, bold=True)
    p = doc.add_paragraph(
        "1. 理解牛顿环干涉现象的形成机理；\n"
        "2. 掌握利用牛顿环测量透镜曲率半径的方法；\n"
        "3. 学习基于图像处理的实验数据自动提取与分析；\n"
        "4. 评估实验结果的精度与误差来源。"
    )
    _set_paragraph_font(p, font_size_pt=11)

    # =====================
    # 二、实验原理
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("二、实验原理"), font_size_pt=14, bold=True)
    p = doc.add_paragraph(
        "牛顿环是由平凸透镜与平玻璃板之间形成的空气薄膜，在单色光照射下产生的等厚干涉条纹。"
        "在反射光条件下，第 n 级暗纹半径 r 满足近似关系：\n"
        "r² = nλR + b。\n"
        "其中 λ 为入射光波长，R 为透镜曲率半径，b 为系统误差项。"
        "通过对 r² 与 n 进行线性拟合，可由斜率求得透镜的曲率半径。"
    )
    _set_paragraph_font(p, font_size_pt=11)

    # =====================
    # 三、实验装置与环境
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("三、实验装置与实验环境"), font_size_pt=14, bold=True)
    p = doc.add_paragraph(
        "牛顿环实验仪、单色光源（钠灯或等效单色光）、相机或电子目镜、计算机，"
        "数据处理软件为 Python（NumPy、OpenCV、Matplotlib 等）。"
    )
    _set_paragraph_font(p, font_size_pt=11)

    # =====================
    # 四、实验方法与数据处理
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("四、实验方法与数据处理"), font_size_pt=14, bold=True)
    p = doc.add_paragraph(
        "首先采集牛顿环干涉图像，并通过图像处理算法确定圆心位置；"
        "随后沿径向提取灰度分布，识别暗纹位置并计算各级暗纹半径；"
        "最终对 r²-n 数据进行线性拟合，得到透镜曲率半径。"
    )
    _set_paragraph_font(p, font_size_pt=11)

    # =====================
    # 五、实验数据与结果
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("五、实验数据与结果"), font_size_pt=14, bold=True)

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "级数 n"
    hdr[1].text = "半径 r / mm"
    hdr[2].text = "r² / mm²"
    for cell in hdr:
        for p in cell.paragraphs:
            _set_paragraph_font(p, font_size_pt=10)

    for n, r, r2 in zip(fit["n"], fit["r_mm"], fit["r2_mm2"]):
        row = table.add_row().cells
        row[0].text = f"{int(n)}"
        row[1].text = f"{r:.4f}"
        row[2].text = f"{r2:.4f}"
        for cell in row:
            for p in cell.paragraphs:
                _set_paragraph_font(p, font_size_pt=10)

    res = doc.add_paragraph()
    run = res.add_run(
        f"线性拟合结果：r² = ({fit['slope']:.6f}) n + ({fit['intercept']:.6f})，"
        f"测得透镜曲率半径 R = {fit['R_mm']:.2f} ± {fit['R_se_mm']:.2f} mm，"
        f"拟合优度 R² = {fit['r_squared']:.4f}。"
    )
    _set_run_font(run, font_size_pt=11)

    # =====================
    # 六、图像结果与分析
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("六、图像结果与分析"), font_size_pt=14, bold=True)

    def add_fig(caption, path):
        if not path or not os.path.exists(path):
            return
        cap = doc.add_paragraph()
        _set_run_font(cap.add_run(caption), font_size_pt=11, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run().add_picture(path, width=Cm(15))

    add_fig("图1 牛顿环暗纹识别叠加图", det.get("overlay_figure"))
    add_fig("图2 径向灰度分布与暗纹位置", det.get("profile_figure"))
    add_fig("图3 r²-n 线性拟合图", fit_fig)

    # =====================
    # 七、误差分析
    # =====================
    if err is not None:
        h = doc.add_paragraph()
        _set_run_font(h.add_run("七、误差分析"), font_size_pt=14, bold=True)

        p = doc.add_paragraph(
            f"拟合优度 R² = {err['r_squared']:.4f}，表明实验数据与理论关系符合较好。"
        )
        _set_paragraph_font(p, font_size_pt=11)

    # =====================
    # 八、结论（最后）
    # =====================
    h = doc.add_paragraph()
    _set_run_font(h.add_run("八、结论"), font_size_pt=14, bold=True)
    p = doc.add_paragraph(
        f"本实验基于图像处理方法对牛顿环干涉条纹进行了分析，"
        f"通过 r²-n 线性拟合测得平凸透镜的曲率半径为 "
        f"R = {fit['R_mm']:.2f} ± {fit['R_se_mm']:.2f} mm。"
        "实验结果与理论模型符合良好，说明该方法具有较高的可行性与测量精度。"
    )
    _set_paragraph_font(p, font_size_pt=11)

    out_path = rep_dir / f"{basename}_NewtonRing_Report.docx"
    doc.save(out_path)

    if logger:
        logger.info("Word report generated: %s", out_path)

    return str(out_path)
